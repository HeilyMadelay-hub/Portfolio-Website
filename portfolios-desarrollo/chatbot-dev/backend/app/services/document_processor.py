import PyPDF2
import docx
from pathlib import Path
from langchain.schema import Document
import logging
from ..config.settings import Config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Procesador para extraer texto de diferentes tipos de documentos"""
    
    def __init__(self):
        self.docs_path = Path(Config.DOCUMENTS_DIR)
        self.supported_extensions = {'.pdf', '.docx', '.md', '.txt'}
    
    def load_all_documents(self):
        """Carga todos los documentos de la carpeta documents/"""
        documents = []
        
        if not self.docs_path.exists():
            logger.warning(f"Directorio de documentos no existe: {self.docs_path}")
            return documents
        
        for file_path in self.docs_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self._process_file(file_path)
                    if doc:
                        documents.append(doc)
                        logger.info(f"Procesado: {file_path.name}")
                except Exception as e:
                    logger.error(f"Error procesando {file_path.name}: {e}")
        
        logger.info(f"Total documentos cargados: {len(documents)}")
        return documents
    
    def _process_file(self, file_path):
        """Procesa un archivo individual según su extensión"""
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._extract_pdf(file_path)
        elif extension == '.docx':
            return self._extract_docx(file_path)
        elif extension in ['.md', '.txt']:
            return self._extract_text(file_path)
        else:
            logger.warning(f"Extensión no soportada: {extension}")
            return None
    
    def _extract_pdf(self, file_path):
        """Extrae texto de archivos PDF"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\\n\\n--- Página {page_num + 1} ---\\n\\n"
                        text += page_text
            
            if text.strip():
                return Document(
                    page_content=text,
                    metadata={
                        'source': str(file_path),
                        'filename': file_path.name,
                        'type': 'pdf',
                        'pages': len(pdf_reader.pages)
                    }
                )
        except Exception as e:
            logger.error(f"Error extrayendo PDF {file_path}: {e}")
        return None
    
    def _extract_docx(self, file_path):
        """Extrae texto de archivos Word"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\\n"
            
            if text.strip():
                return Document(
                    page_content=text,
                    metadata={
                        'source': str(file_path),
                        'filename': file_path.name,
                        'type': 'docx'
                    }
                )
        except Exception as e:
            logger.error(f"Error extrayendo DOCX {file_path}: {e}")
        return None
    
    def _extract_text(self, file_path):
        """Extrae texto de archivos de texto plano y markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if text.strip():
                return Document(
                    page_content=text,
                    metadata={
                        'source': str(file_path),
                        'filename': file_path.name,
                        'type': file_path.suffix[1:]  # sin el punto
                    }
                )
        except Exception as e:
            logger.error(f"Error extrayendo texto {file_path}: {e}")
        return None
    
    def get_documents_info(self):
        """Retorna información sobre los documentos disponibles"""
        info = {
            'total_files': 0,
            'by_type': {},
            'files': []
        }
        
        if not self.docs_path.exists():
            return info
        
        for file_path in self.docs_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                file_type = file_path.suffix.lower()[1:]  # sin el punto
                
                info['total_files'] += 1
                info['by_type'][file_type] = info['by_type'].get(file_type, 0) + 1
                info['files'].append({
                    'name': file_path.name,
                    'type': file_type,
                    'size': file_path.stat().st_size
                })
        
        return info